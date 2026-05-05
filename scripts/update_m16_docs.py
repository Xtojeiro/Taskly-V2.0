from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "Envio-27.01"
DESIGN = ROOT / "focus. Design"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(13 if level == 1 else 12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_center(doc: Document, text: str, size: int = 11) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_text(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    doc.add_paragraph()


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.55)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    return doc


def save_proposal() -> Path:
    doc = setup_doc()
    add_center(doc, "PROPOSTA DE PROJETO (SOFTWARE) - M16", 13)
    add_center(doc, 'CURSO PROFISSIONAL "TECNICO DE GESTAO E PROGRAMACAO DE SISTEMAS INFORMATICOS"')
    add_center(doc, "TRIENIO 2025/2026")
    add_center(doc, "NOME DO ALUNO: Tomas Xavier Marques Tojeiro    TURMA: 12E    No 23")

    add_heading(doc, "TEMA E TIPOLOGIA DO PROJETO M16", 2)
    add_text(
        doc,
        "O projeto chama-se Taskly e consiste numa aplicacao mobile de rastreamento de tarefas e habitos, "
        "desenvolvida com Expo/React Native e Convex. A aplicacao ajuda o utilizador a organizar rotinas, "
        "acompanhar a consistencia diaria e receber sugestoes personalizadas com inteligencia artificial via OpenRouter.",
    )

    add_heading(doc, "OBJETIVOS DO TEMA/PROJETO", 2)
    add_bullets(
        doc,
        [
            "Permitir registo e login simples para separar os dados de cada utilizador.",
            "Criar, consultar e concluir tarefas/habitos recorrentes diarios, semanais ou mensais.",
            "Apresentar progresso diario, pontuacao e resumo de tarefas concluidas.",
            "Disponibilizar um heatmap mensal interativo para visualizar consistencia ao longo do tempo.",
            "Permitir personalizacao visual atraves de cor de foco e preferencia de tema.",
            "Gerar sugestoes praticas com IA, mantendo a chave OpenRouter apenas no backend Convex.",
            "Documentar o modelo de dados, requisitos, limites e fluxo de utilizacao da aplicacao.",
        ],
    )

    add_heading(doc, "FASES DO DESENVOLVIMENTO DO PROJETO (DATAS)", 2)
    add_table(
        doc,
        ["Fase", "Periodo", "Resultado esperado"],
        [
            ["1. Planeamento e proposta", "Janeiro 2026", "Definicao do tema, objetivos, tecnologias e recursos."],
            ["2. Analise do sistema", "Janeiro/Fevereiro 2026", "Requisitos funcionais, modelo de dados e prototipos."],
            ["3. Implementacao base", "Marco/Abril 2026", "App Expo com navegacao, Convex, auth local e CRUD."],
            ["4. Estatisticas e IA", "Abril/Maio 2026", "Heatmap, pontuacao, estatisticas e insights OpenRouter."],
            ["5. Testes e entrega", "Maio 2026", "Correcao de erros, testes, lint/typecheck e documentacao final."],
        ],
    )

    add_heading(doc, "RECURSOS PREVISTOS E METODOLOGIAS A UTILIZAR NA CONCRETIZACAO DO PROJETO", 2)
    add_bullets(
        doc,
        [
            "Computador portatil com acesso a internet para desenvolvimento e testes.",
            "Visual Studio Code Insiders como IDE principal.",
            "Linha de comandos Windows para dependencias, Convex, Expo, testes e diagnostico de erros.",
            "Expo Go para execucao e validacao em telemovel.",
            "Convex como backend, base de dados reativa, mutations, queries e actions.",
            "OpenRouter para integracao de inteligencia artificial via API key guardada no ambiente Convex.",
            "Metodologia incremental: implementar uma funcionalidade, testar, corrigir, documentar e repetir.",
            "Validacao com npx tsc --noEmit, npm run lint e npm test.",
        ],
    )

    add_heading(doc, "DATA DE RECECAO", 2)
    add_text(doc, "_____/_____/________                                      A Professora")
    add_text(doc, "_____________________________                 (Ana Mafalda Silva)")
    output = DOCS / "0_PROPOSTA_PROJETO_M16_25_26.docx"
    doc.save(output)
    return output


def save_analysis() -> Path:
    doc = setup_doc()
    add_center(doc, "ANALISE DO SISTEMA PARA O PROJETO DE SOFTWARE - M16", 13)
    add_center(doc, 'CURSO PROFISSIONAL "TECNICO DE GESTAO E PROGRAMACAO DE SISTEMAS INFORMATICOS"')
    add_center(doc, "TRIENIO 2025/2026")
    add_center(doc, "NOME DO ALUNO: Tomas Tojeiro    TURMA: 12E    No 23")

    add_heading(doc, "TEMA E TIPOLOGIA DO PROJETO M16", 2)
    add_text(
        doc,
        "Taskly e uma aplicacao mobile de produtividade e acompanhamento de habitos. O produto final combina "
        "organizacao diaria, gamificacao atraves de pontos, visualizacao por heatmap e sugestoes de melhoria geradas por IA.",
    )

    add_heading(doc, "REQUISITOS FUNCIONAIS", 2)
    add_bullets(
        doc,
        [
            "O utilizador pode criar conta e iniciar sessao com email e password num sistema de autenticacao local simples.",
            "A app guarda um token de sessao no dispositivo, permitindo retomar a utilizacao sem novo login.",
            "O utilizador pode criar tarefas/habitos com nome, categoria, frequencia, dias alvo, hora e descricao.",
            "A app apresenta uma vista Today com progresso, pontos, tarefas do dia e cartao de AI Insights.",
            "A vista To-Do permite alternar entre Daily, Weekly e Range, ordenar por hora e marcar/desmarcar tarefas.",
            "A vista Stats apresenta heatmap mensal interativo, percentagem de conclusao, pontos e detalhes do dia selecionado.",
            "A vista Settings permite alterar cor de foco, preferencia de tema e terminar sessao.",
            "Ao concluir tarefas, o backend recalcula estatisticas diarias e atualiza pontuacao.",
            "A action de IA envia um resumo dos dados para OpenRouter e guarda 2 a 4 sugestoes praticas.",
        ],
    )

    add_heading(doc, "MODELO DE DADOS", 2)
    add_text(doc, "O backend utiliza Convex. As tabelas principais sao:")
    add_table(
        doc,
        ["Tabela", "Campos principais", "Relacoes/Indices"],
        [
            ["users", "_id, name, email, tokenIdentifier, passwordHash, goals, targetPerWeek, focusColor, theme", "by_email, by_token"],
            ["sessions", "_id, userId, token, createdAt, expiresAt", "sessions pertence a users; by_token, by_userId"],
            ["tags", "_id, userId, name, color, createdAt", "tags pertence a users; by_userId_and_name"],
            ["habits", "_id, userId, name, description, category, frequency, targetDays, scheduledTime, isActive", "habits pertence a users; by_userId_and_isActive"],
            ["habitLogs", "_id, habitId, userId, date, completed, completionRate", "logs pertence a habits e users; by_habitId_and_date, by_userId_and_date"],
            ["dailyStats", "_id, userId, date, totalHabits, completedHabits, completionRate, score", "estatisticas por utilizador e data; by_userId_and_date"],
            ["aiInsights", "_id, userId, summary, suggestions, model, createdAt", "insights pertence a users; by_userId_and_createdAt"],
        ],
    )

    add_heading(doc, "INTEGRACAO COM OPENROUTER", 2)
    add_text(
        doc,
        "A integracao de IA e feita no backend Convex por uma action chamada generateInsights. A chave "
        "OPENROUTER_API_KEY fica apenas nas variaveis de ambiente Convex. O modelo por defeito e openrouter/auto, "
        "permitindo ao OpenRouter escolher automaticamente um modelo adequado.",
    )

    add_heading(doc, "LIMITACOES E SEGURANCA", 2)
    add_bullets(
        doc,
        [
            "A autenticacao e local e simples, adequada para demonstracao escolar, nao para producao.",
            "A password e protegida por hash simples, mas sem todos os mecanismos de seguranca de uma solucao profissional.",
            "A API key OpenRouter nao deve ser colocada em EXPO_PUBLIC_* nem no codigo frontend.",
            "O projeto e mobile-first; a execucao web e secundaria.",
        ],
    )

    add_heading(doc, "INTERFACE DO SOFTWARE", 2)
    add_text(doc, "Ecras de referencia usados no desenvolvimento:")
    for image_name in [
        "login-screen.png",
        "Main-screen.png",
        "todo-screen.png",
        "interactive-heatmap.png",
        "settings.png",
    ]:
        image_path = DESIGN / image_name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(2.0))

    add_heading(doc, "DATA DE RECECAO", 2)
    add_text(doc, "_____/_____/________                                      A Professora")
    add_text(doc, "_____________________________                 (Ana Mafalda Silva)")
    output = DOCS / "1_Analise_do_sistema_PROJETO_SOFTWARE_M16_Taskly.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    print(save_proposal())
    print(save_analysis())
